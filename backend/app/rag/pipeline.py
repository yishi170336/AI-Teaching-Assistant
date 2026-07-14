from __future__ import annotations

import hashlib
import json
import logging
import re
import threading
import unicodedata
from collections import Counter
from dataclasses import replace
from pathlib import Path
from typing import Any, Callable, Iterable

import fitz
import numpy as np
from docx import Document
from openpyxl import load_workbook

from backend.app.config import settings
from backend.app.rag.models import PageDocument, TextChunk
from backend.app.rag.embedding_runtime import encode_texts
from backend.app.rag.ontology import (
    COURSE_CONCEPTS,
    extract_course_concepts,
    is_course_concept,
    normalize_concept_name,
)
from backend.app.rag.multimodal import (
    BuildModelConfig,
    LayoutElement,
    SCANNED_PAGE_PLACEHOLDER,
    build_local_knowledge_graph,
    enhance_pdf,
    multimodal_chunks,
)
from backend.app.rag.stores import build_qdrant_indexes, sync_neo4j_graph


logger = logging.getLogger(__name__)


class KnowledgeBaseBuildCancelled(RuntimeError):
    """Raised when a cooperative knowledge-base build cancellation is requested."""


BuildProgressCallback = Callable[[int, str, str], None]

KNOWLEDGE_EXTENSIONS = {".pdf", ".md", ".txt", ".docx"}
QUESTION_BANK_EXTENSIONS = {".xlsx", ".json"}
SUPPORTED_EXTENSIONS = KNOWLEDGE_EXTENSIONS | QUESTION_BANK_EXTENSIONS
AD_NOISE = (
    "扫码关注",
    "微信公众号",
    "关注公众号",
    "购买正版",
    "资源下载",
    "广告",
)
TAG_KEYWORDS = COURSE_CONCEPTS


def _normalize_line(line: str) -> str:
    line = unicodedata.normalize("NFKC", line).replace("\u200b", "")
    line = re.sub(r"https?\s*:\s*[/\\]+\s*\S+", "", line, flags=re.I)
    line = re.sub(r"www\s*\.\s*\S+", "", line, flags=re.I)
    line = re.sub(r"(?<=[A-Za-z])\s+(?=[A-Za-z0-9])", "", line)
    line = re.sub(r"(?<=[0-9])\s+(?=[A-Za-z])", "", line)
    line = re.sub(r"(?<=[A-Za-z])\s+(?=[0-9])", "", line)
    line = re.sub(r"P\s*N\s*结", "PN结", line, flags=re.I)
    line = re.sub(r"([NP])\s*型", r"\1型", line, flags=re.I)
    line = re.sub(r"M\s*O\s*S", "MOS", line, flags=re.I)
    line = re.sub(r"[ \t]+", " ", line).strip()
    return line


def _edge_noise(raw_pages: list[str]) -> set[str]:
    candidates: Counter[str] = Counter()
    for text in raw_pages:
        lines = [_normalize_line(line) for line in text.splitlines() if _normalize_line(line)]
        for line in lines[:2] + lines[-2:]:
            if len(line) <= 80:
                candidates[line] += 1
    threshold = max(3, int(len(raw_pages) * 0.08))
    return {line for line, count in candidates.items() if count >= threshold}


def clean_page_text(text: str, repeated_noise: set[str] | None = None) -> str:
    repeated_noise = repeated_noise or set()
    cleaned_lines: list[str] = []
    for raw_line in text.splitlines():
        line = _normalize_line(raw_line)
        if not line or line in repeated_noise:
            continue
        if re.fullmatch(r"[-—·•\s]*\d{1,4}[-—·•\s]*", line):
            continue
        if any(noise in line for noise in AD_NOISE):
            continue
        if sum(char == "�" for char in line) > 1:
            continue
        cleaned_lines.append(line)

    paragraphs: list[str] = []
    buffer = ""
    heading_pattern = re.compile(
        r"^(?:第[一二三四五六七八九十百0-9]+章|\d+(?:\.\d+){0,3}\s+|本章小结|自测题|习题)"
    )
    for line in cleaned_lines:
        is_heading = bool(heading_pattern.match(line)) and len(line) < 70
        if is_heading:
            if buffer:
                paragraphs.append(buffer)
                buffer = ""
            paragraphs.append(line)
            continue
        buffer += line
        if line.endswith(("。", "！", "？", ":", "；")) or len(buffer) >= 260:
            paragraphs.append(buffer)
            buffer = ""
    if buffer:
        paragraphs.append(buffer)
    return "\n\n".join(paragraphs)


def _is_chapter_title(title: str) -> bool:
    return bool(re.match(r"^第[一二三四五六七八九十百0-9]+章", title.replace(" ", "")))


def extract_pdf(path: Path, chapter_limit: int | None = None) -> list[PageDocument]:
    document = fitz.open(path)
    toc = [item for item in document.get_toc(simple=True) if len(item) >= 3]
    chapters = [(str(title).strip(), int(page)) for level, title, page in toc if level == 1 and _is_chapter_title(str(title).strip())]
    if chapters:
        start_page = chapters[0][1]
        if chapter_limit and len(chapters) > chapter_limit:
            end_page = chapters[chapter_limit][1] - 1
        else:
            end_page = document.page_count
    else:
        start_page, end_page = 1, document.page_count

    raw_pages = [document[index - 1].get_text("text") for index in range(start_page, end_page + 1)]
    repeated_noise = _edge_noise(raw_pages)
    page_docs: list[PageDocument] = []
    page_range_match = re.search(r"(?:pages?|页)[_-]?(\d+)[_-](\d+)", path.stem, re.I)
    source_page_offset = int(page_range_match.group(1)) - 1 if page_range_match else 0
    for page_number, raw_text in zip(range(start_page, end_page + 1), raw_pages):
        current_chapter = ""
        current_section = ""
        for level, title, toc_page in toc:
            if int(toc_page) > page_number:
                break
            title = _normalize_line(re.sub(r"\s+", " ", str(title)).strip())
            if level == 1 and _is_chapter_title(title):
                current_chapter = title
                current_section = ""
            elif level >= 2:
                current_section = title
        text = clean_page_text(raw_text, repeated_noise)
        page_object = document[page_number - 1]
        has_visual_content = bool(page_object.get_images(full=True)) or len(page_object.get_drawings()) >= 3
        has_formula_content = bool(
            re.search(r"[=+−±√∫ΣΩπ^_].*[A-Za-z0-9]|[A-Za-z0-9].*[=+−±√∫ΣΩπ^_]", text)
        )
        if len(text) < 30 and not has_visual_content and not has_formula_content:
            continue
        if not text:
            text = SCANNED_PAGE_PLACEHOLDER
        page_docs.append(
            PageDocument(
                text=text,
                source=path.name,
                page=page_number,
                chapter=current_chapter or path.stem,
                section=current_section or current_chapter or path.stem,
                source_page=source_page_offset + page_number,
            )
        )
    document.close()
    return page_docs


def extract_markdown_or_text(path: Path) -> list[PageDocument]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    chapter = path.stem
    section = path.stem
    documents: list[PageDocument] = []
    page = 1
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        cleaned = clean_page_text("\n".join(buffer))
        if cleaned:
            documents.append(PageDocument(cleaned, path.name, page, chapter, section))
        buffer = []

    for line in text.splitlines():
        heading = re.match(r"^(#{1,3})\s+(.+)$", line.strip())
        if heading:
            flush()
            title = heading.group(2).strip()
            if len(heading.group(1)) == 1:
                chapter = title
            section = title
        elif line.strip() == "\f":
            flush()
            page += 1
        else:
            buffer.append(line)
    flush()
    return documents


def extract_docx(path: Path) -> list[PageDocument]:
    document = Document(path)
    chapter = path.stem
    section = path.stem
    blocks: list[PageDocument] = []
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        text = clean_page_text("\n".join(buffer))
        if text:
            blocks.append(PageDocument(text, path.name, 1, chapter, section))
        buffer = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if "heading" in style or "标题" in style:
            flush()
            section = text
            if style.endswith("1"):
                chapter = text
        else:
            buffer.append(text)
    flush()
    return blocks


def _split_tags(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return [item.strip() for item in re.split(r"[,，、;；|]", str(value or "")) if item.strip()]


QUESTION_HEADER_ALIASES = {
    "question_id": ("题号", "question_id", "id"),
    "question_text": ("题目文本", "题目", "question_text"),
    "knowledge_tags": ("知识点标签", "知识点", "knowledge_tags"),
    "standard_answer": ("标准答案", "答案", "standard_answer"),
    "common_mistakes": ("易错点", "common_mistakes"),
    "difficulty": ("难度", "difficulty"),
    "question_type": ("题型", "question_type"),
    "solution_steps": ("解题步骤", "解析", "solution_steps"),
}


def extract_question_xlsx(path: Path) -> list[dict[str, Any]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    sheet = workbook.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return []
    header_index = next(
        (
            index
            for index, row in enumerate(rows[:12])
            if "题号" in [str(value or "").strip() for value in row]
            and "题目文本" in [str(value or "").strip() for value in row]
        ),
        0,
    )
    headers = [str(value or "").strip() for value in rows[header_index]]
    mapping: dict[str, int] = {}
    for field, aliases in QUESTION_HEADER_ALIASES.items():
        for alias in aliases:
            if alias in headers:
                mapping[field] = headers.index(alias)
                break
    required = {"question_id", "question_text", "knowledge_tags", "standard_answer", "common_mistakes"}
    if not required.issubset(mapping):
        workbook.close()
        return []
    questions: list[dict[str, Any]] = []
    for row in rows[header_index + 1 :]:
        if not row or not row[mapping["question_text"]]:
            continue
        item = {
            field: (row[index] if index < len(row) else "")
            for field, index in mapping.items()
        }
        item["question_id"] = str(item["question_id"])
        item["question_text"] = str(item["question_text"]).strip()
        item["knowledge_tags"] = _split_tags(item["knowledge_tags"])
        item["standard_answer"] = str(item["standard_answer"] or "").strip()
        item["common_mistakes"] = str(item["common_mistakes"] or "").strip()
        item["difficulty"] = str(item.get("difficulty") or "基础").strip()
        item["question_type"] = str(item.get("question_type") or "综合题").strip()
        item["solution_steps"] = str(item.get("solution_steps") or "").strip()
        item["source"] = path.name
        questions.append(item)
    workbook.close()
    return questions


def extract_question_json(path: Path) -> list[dict[str, Any]]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, UnicodeError):
        return []
    if isinstance(data, dict):
        data = data.get("questions", [])
    if not isinstance(data, list):
        return []
    questions = []
    for index, item in enumerate(data, 1):
        if not isinstance(item, dict) or not item.get("question_text"):
            continue
        normalized = dict(item)
        normalized.setdefault("question_id", f"JSON-{index:03d}")
        normalized["knowledge_tags"] = _split_tags(normalized.get("knowledge_tags"))
        normalized.setdefault("standard_answer", "")
        normalized.setdefault("common_mistakes", "")
        normalized.setdefault("difficulty", "基础")
        normalized.setdefault("question_type", "综合题")
        normalized.setdefault("solution_steps", "")
        normalized["source"] = path.name
        questions.append(normalized)
    return questions


def _knowledge_tags(
    text: str,
    section: str,
    supplemental: Iterable[str] = (),
) -> list[str]:
    tags = extract_course_concepts(text, section)
    compact_text = re.sub(r"\s+", "", text).lower()
    for value in supplemental:
        concept = normalize_concept_name(str(value))
        if (
            concept
            and is_course_concept(concept)
            and re.sub(r"\s+", "", concept).lower() in compact_text
            and concept not in tags
        ):
            tags.append(concept)
    return tags[:12]


def _sentence_pieces(text: str, max_chars: int = 900) -> list[str]:
    paragraphs = [paragraph.strip() for paragraph in text.split("\n\n") if paragraph.strip()]
    pieces: list[str] = []
    for paragraph in paragraphs:
        if len(paragraph) <= max_chars:
            pieces.append(paragraph)
            continue
        sentences = [piece for piece in re.split(r"(?<=[。！？；])", paragraph) if piece]
        current = ""
        for sentence in sentences:
            if current and len(current) + len(sentence) > max_chars:
                pieces.append(current)
                current = ""
            if len(sentence) > max_chars:
                pieces.extend(
                    sentence[index : index + max_chars]
                    for index in range(0, len(sentence), max_chars)
                )
                continue
            current += sentence
        if current:
            pieces.append(current)
    return pieces


def chunk_documents(documents: Iterable[PageDocument], max_chars: int = 900) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    for document in documents:
        pieces = _sentence_pieces(document.text, max_chars=max_chars)
        supplemental_concepts = (
            document.extra.get("ocr_concepts", [])
            if isinstance(document.extra, dict)
            and isinstance(document.extra.get("ocr_concepts"), list)
            else []
        )
        current: list[str] = []
        current_length = 0

        def flush() -> None:
            nonlocal current, current_length
            if not current:
                return
            text = "\n\n".join(current).strip()
            raw_id = f"{document.source}|{document.page}|{document.section}|{text[:120]}"
            chunk_id = hashlib.sha1(raw_id.encode("utf-8")).hexdigest()[:16]
            chunks.append(
                TextChunk(
                    id=chunk_id,
                    text=text,
                    source=document.source,
                    chapter=document.chapter,
                    section=document.section,
                    page_start=document.source_page or document.page,
                    page_end=document.source_page or document.page,
                    doc_type=document.doc_type,
                    knowledge_tags=_knowledge_tags(
                        text, document.section, supplemental_concepts
                    ),
                    element_type=document.element_type,
                    bbox=document.bbox,
                    parent_id=document.parent_id,
                    image_path=document.image_path,
                    content_hash=document.content_hash,
                    multimodal=document.extra,
                )
            )
            overlap = text[-140:] if len(text) > 140 else ""
            current = [overlap] if overlap else []
            current_length = len(overlap)

        for piece in pieces:
            if current and current_length + len(piece) > max_chars:
                flush()
            if current and current_length + len(piece) > max_chars:
                current = []
                current_length = 0
            current.append(piece)
            current_length += len(piece)
        flush()
    return chunks


def question_chunks(questions: Iterable[dict[str, Any]]) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    for item in questions:
        tags = _split_tags(item.get("knowledge_tags"))
        text = (
            f"题目：{item.get('question_text', '')}\n"
            f"标准答案：{item.get('standard_answer', '')}\n"
            f"解题步骤：{item.get('solution_steps', '')}\n"
            f"易错点：{item.get('common_mistakes', '')}"
        ).strip()
        question_id = str(item.get("question_id", "Q"))
        chunks.append(
            TextChunk(
                id=f"question-{question_id}",
                text=text,
                source=str(item.get("source", "question_bank.json")),
                chapter="示例题库",
                section="、".join(tags) or "综合",
                page_start=None,
                page_end=None,
                doc_type="question",
                knowledge_tags=tags,
            )
        )
    return chunks


def _write_clean_markdown(path: Path, documents: list[PageDocument], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    lines = [f"# {path.stem}", "", f"> 清洗来源：{path.name}", ""]
    last_chapter = last_section = ""
    for document in documents:
        if document.chapter and document.chapter != last_chapter:
            lines.extend([f"## {document.chapter}", ""])
            last_chapter = document.chapter
        if document.section and document.section != last_section and document.section != document.chapter:
            lines.extend([f"### {document.section}", ""])
            last_section = document.section
        lines.extend([
            f"<!-- source={document.source}; page={document.source_page or document.page} -->",
            document.text,
            "",
        ])
    (output_dir / f"{path.stem}.clean.md").write_text("\n".join(lines), encoding="utf-8")


def validate_extracted_content(chunks: list[TextChunk]) -> dict[str, int]:
    text_chunks = [
        chunk for chunk in chunks
        if chunk.doc_type == "textbook" and chunk.element_type == "text"
    ]
    placeholders = [
        chunk for chunk in text_chunks
        if chunk.text.strip() == SCANNED_PAGE_PLACEHOLDER
    ]
    if placeholders and len(placeholders) >= max(3, (len(text_chunks) + 3) // 4):
        raise RuntimeError(
            "扫描版 PDF 的正文 OCR 未成功："
            f"{len(placeholders)}/{len(text_chunks)} 个正文片段仍是图形占位符。"
            "请配置可用的 Qwen3-VL API 后重建，旧索引不会被替换。"
        )
    return {
        "text_chunks": len(text_chunks),
        "placeholder_text_chunks": len(placeholders),
    }


def validate_graph_semantics(
    chunks: list[TextChunk], graph: dict[str, Any]
) -> dict[str, int]:
    textbook_chunks = [chunk for chunk in chunks if chunk.doc_type == "textbook"]
    concepts = {
        str(node.get("name", "")).strip()
        for node in graph.get("nodes", [])
        if node.get("type") == "concept" and str(node.get("name", "")).strip()
    }
    if len(textbook_chunks) >= 12 and len(concepts) <= 1:
        raise RuntimeError(
            "知识图谱语义校验失败：教材内容较多，但只识别出"
            f" {len(concepts)} 个知识点。请检查 OCR/章节识别结果，旧索引不会被替换。"
        )
    return {"concept_nodes": len(concepts)}


def validate_build_artifacts(
    chunks: list[TextChunk],
    embeddings: np.ndarray,
    graph: dict[str, Any],
) -> dict[str, Any]:
    if any(chunk.doc_type == "question" for chunk in chunks):
        raise RuntimeError("构建校验失败：题库 Chunk 不得进入课程知识库")
    if embeddings.ndim != 2 or embeddings.shape[0] != len(chunks):
        raise RuntimeError("构建校验失败：向量数量与 Chunk 数量不一致")
    node_ids = {
        str(node.get("id", "")) for node in graph.get("nodes", []) if node.get("id")
    }
    dangling = [
        edge for edge in graph.get("edges", [])
        if str(edge.get("source", "")) not in node_ids
        or str(edge.get("target", "")) not in node_ids
    ]
    if dangling:
        raise RuntimeError(f"构建校验失败：知识图谱存在 {len(dangling)} 条悬空关系")
    invalid_bbox = [
        chunk.id for chunk in chunks
        if chunk.bbox is not None and len(chunk.bbox) != 4
    ]
    if invalid_bbox:
        raise RuntimeError(f"构建校验失败：{len(invalid_bbox)} 个多模态元素坐标不完整")
    return {
        "status": "passed",
        "chunks": len(chunks),
        "vectors": int(embeddings.shape[0]),
        "vector_dimension": int(embeddings.shape[1]),
        "graph_nodes": len(node_ids),
        "graph_edges": len(graph.get("edges", [])),
        "question_chunks": 0,
        "dangling_graph_edges": 0,
        "concept_nodes": sum(
            node.get("type") == "concept" for node in graph.get("nodes", [])
        ),
        "placeholder_text_chunks": sum(
            chunk.text.strip() == SCANNED_PAGE_PLACEHOLDER
            for chunk in chunks
            if chunk.doc_type == "textbook" and chunk.element_type == "text"
        ),
    }


def _formula_pipeline_stats(output_dir: Path, elements: list[LayoutElement]) -> dict[str, Any]:
    categories: Counter[str] = Counter()
    audit_counts: Counter[str] = Counter()
    for path in output_dir.glob("*.pdf_extract_kit.json"):
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        for region in payload.get("regions", []):
            if not isinstance(region, dict):
                continue
            category = str(region.get("category", "")).lower()
            detector = str(region.get("detector", ""))
            if detector.endswith(":formula") or category in {
                "isolate_formula", "isolated", "isolated_formula"
            }:
                categories[category] += 1
    for path in output_dir.glob("*.formula_audit.json"):
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        for name in ("detected", "recognized", "fallback", "uncertain"):
            try:
                audit_counts[name] += int(payload.get(name, 0))
            except (TypeError, ValueError):
                continue
    formulas = [element for element in elements if element.element_type == "formula"]
    display_candidates = audit_counts["detected"] or sum(
        categories.get(name, 0)
        for name in ("isolate_formula", "isolated", "isolated_formula")
    )
    return {
        "detected_regions": sum(categories.values()),
        "inline_regions_kept_in_text": categories.get("inline", 0),
        "display_candidates": display_candidates,
        "recognized_formulas": audit_counts["recognized"],
        "fallback_formulas": audit_counts["fallback"],
        "indexed_formulas": len(formulas),
        "rejected_or_merged_regions": max(0, display_candidates - len(formulas)),
        "uncertain_formulas": (
            audit_counts["uncertain"]
            if audit_counts["detected"]
            else sum(element.uncertain for element in formulas)
        ),
        "recognition": (
            "PDF-Extract-Kit localization + native PDF geometry LaTeX + "
            f"qwen/{settings.qwen_circuit_vision_model} scan fallback"
            if settings.qwen_api_key
            else "PDF-Extract-Kit localization + PyMuPDF text fallback"
        ),
    }


def build_knowledge_base(
    resources_dir: Path,
    output_dir: Path,
    embedding_model_path: Path,
    *,
    chapter_limit: int | None = None,
    model_config: BuildModelConfig | None = None,
    knowledge_base_id: str | None = None,
    sync_graph_store: bool = True,
    progress_callback: BuildProgressCallback | None = None,
    cancel_event: threading.Event | None = None,
) -> dict[str, Any]:
    def report(progress: int, stage: str, message: str) -> None:
        if progress_callback is not None:
            progress_callback(progress, stage, message)
        if cancel_event is not None and cancel_event.is_set():
            raise KnowledgeBaseBuildCancelled("用户已取消知识库构建")

    resources_dir = resources_dir.resolve()
    output_dir = output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    cleaned_dir = output_dir / "cleaned_documents"
    cleaned_dir.mkdir(parents=True, exist_ok=True)
    report(5, "document_scanning", "正在扫描教材与讲义")

    documents: list[PageDocument] = []
    elements: list[LayoutElement] = []
    cleaning_audits: list[dict[str, Any]] = []
    candidate_files = [
        path for path in sorted(resources_dir.iterdir())
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    source_files = [
        path for path in candidate_files if path.suffix.lower() in KNOWLEDGE_EXTENSIONS
    ]
    excluded_sources = [
        {
            "source": path.name,
            "reason": "题库文件与课程知识库隔离，不参与分块、向量、图谱或检索",
        }
        for path in candidate_files
        if path.suffix.lower() in QUESTION_BANK_EXTENSIONS
    ]
    source_count = max(1, len(source_files))
    for source_index, path in enumerate(source_files):
        report(
            10 + int(source_index / source_count * 35),
            "document_parsing",
            f"正在解析 {path.name}（{source_index + 1}/{len(source_files)}）",
        )
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            extracted = extract_pdf(path, chapter_limit)
            extracted, pdf_elements, audit = enhance_pdf(
                path,
                extracted,
                output_dir,
                model_config=model_config,
            )
            repeated_noise = _edge_noise([item.text for item in extracted])
            extracted = [
                replace(
                    item,
                    text=clean_page_text(item.text, repeated_noise) or item.text,
                )
                for item in extracted
            ]
            documents.extend(extracted)
            elements.extend(pdf_elements)
            cleaning_audits.extend(
                {**item, "source": path.name} for item in audit
            )
            _write_clean_markdown(path, extracted, cleaned_dir)
        elif suffix in {".md", ".txt"}:
            extracted = extract_markdown_or_text(path)
            documents.extend(extracted)
            _write_clean_markdown(path, extracted, cleaned_dir)
        elif suffix == ".docx":
            extracted = extract_docx(path)
            documents.extend(extracted)
            _write_clean_markdown(path, extracted, cleaned_dir)
        report(
            10 + int((source_index + 1) / source_count * 35),
            "document_cleaning",
            f"已完成 {path.name} 的解析与清洗",
        )
    structured_questions = {
        "schema_version": "2.0",
        "questions": [],
        "excluded_sources": excluded_sources,
        "message": "题库与课程知识库隔离；本文件仅保留兼容占位。",
    }
    (output_dir / "question_bank.json").write_text(
        json.dumps(structured_questions, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    report(50, "chunking", "正在切分文本并整理多模态元素")
    chunks = chunk_documents(documents) + multimodal_chunks(elements)
    if not chunks:
        raise RuntimeError(f"在 {resources_dir} 中没有提取到可索引内容")
    extraction_quality = validate_extracted_content(chunks)

    chunk_path = output_dir / "chunks.jsonl"
    chunk_path.write_text(
        "\n".join(json.dumps(chunk.to_dict(), ensure_ascii=False) for chunk in chunks),
        encoding="utf-8",
    )
    (output_dir / "multimodal_elements.jsonl").write_text(
        "\n".join(json.dumps(element.to_dict(), ensure_ascii=False) for element in elements),
        encoding="utf-8",
    )
    (output_dir / "cleaning_audit.json").write_text(
        json.dumps(cleaning_audits, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    report(60, "knowledge_graph", "正在构建知识图谱关系")
    graph = build_local_knowledge_graph(chunks)
    semantic_quality = validate_graph_semantics(chunks, graph)
    (output_dir / "knowledge_graph.json").write_text(
        json.dumps(graph, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    if settings.qdrant_url:
        # Establish native-module import order before Torch on Windows.
        import qdrant_client  # noqa: F401

    embedding_texts = [
        "\n".join(
            filter(
                None,
                [chunk.doc_type, chunk.chapter, chunk.section, " ".join(chunk.knowledge_tags), chunk.text],
            )
        )
        for chunk in chunks
    ]
    report(68, "embedding", f"正在生成 {len(chunks)} 个内容向量")
    embeddings = encode_texts(
        embedding_model_path,
        embedding_texts,
        batch_size=32,
        show_progress_bar=True,
    )
    report(82, "validation", "正在校验向量与图谱完整性")
    validation = validate_build_artifacts(chunks, embeddings, graph)
    import faiss

    index = faiss.IndexFlatIP(embeddings.shape[1])
    index.add(embeddings)
    # FAISS' Windows file writer cannot open paths containing Chinese characters.
    # Serialize in memory and let Python handle the Unicode path instead.
    serialized_index = faiss.serialize_index(index)
    (output_dir / "vectors.faiss").write_bytes(serialized_index.tobytes())

    report(85, "indexing", "正在写入向量索引")
    qdrant_status = build_qdrant_indexes(output_dir, chunks, embeddings)
    neo4j_status = (
        sync_neo4j_graph(knowledge_base_id or output_dir.name, graph)
        if sync_graph_store
        else {"enabled": False, "reason": "deferred until atomic index activation"}
    )

    formula_processing = _formula_pipeline_stats(output_dir, elements)
    metadata = {
        "state": "populated",
        "schema_version": "2.3-circuit-image-retrieval",
        "resource_dir": str(resources_dir),
        "embedding_model": str(embedding_model_path),
        "dimension": int(embeddings.shape[1]),
        "documents": len(source_files),
        "text_pages": len(documents),
        "ocr_pages": sum(
            isinstance(item.extra, dict) and bool(item.extra.get("ocr_processor"))
            for item in documents
        ),
        "questions": 0,
        "excluded_sources": excluded_sources,
        "chunks": len(chunks),
        "layout_elements": len(elements),
        "circuit_diagrams": sum(item.element_type == "circuit" for item in elements),
        "formula_elements": sum(item.element_type == "formula" for item in elements),
        "formula_processing": formula_processing,
        "table_elements": sum(item.element_type == "table" for item in elements),
        "discarded_pages": sum(not item.get("keep", True) for item in cleaning_audits),
        "knowledge_graph": {"nodes": len(graph["nodes"]), "edges": len(graph["edges"]), "neo4j": neo4j_status},
        "qdrant": qdrant_status,
        "vision_model": (
            f"qwen/{settings.qwen_circuit_vision_model}"
            if settings.qwen_api_key
            else "not-configured (safe fallback)"
        ),
        "pdf_extract_kit": (
            json.loads((output_dir / "pdf_extract_kit_manifest.json").read_text(encoding="utf-8"))
            if (output_dir / "pdf_extract_kit_manifest.json").exists()
            else {"enabled": False}
        ),
        "chapter_limit": chapter_limit,
        "sources": [path.name for path in source_files],
        "validation": validation,
        "extraction_quality": {**extraction_quality, **semantic_quality},
    }
    metadata["pipeline_layers"] = {
        "document_cleaning": {
            "status": "ready",
            "pages_reviewed": len(cleaning_audits),
            "pages_discarded": metadata["discarded_pages"],
            "partial_characters_removed": sum(
                int(item.get("removed_characters", 0)) for item in cleaning_audits
            ),
            "question_banks_excluded": len(excluded_sources),
        },
        "document_parsing": {
            "status": "ready",
            "engine": "PyMuPDF text + Qwen3-VL scanned-page OCR + PDF-Extract-Kit layout",
            "ocr_pages": metadata["ocr_pages"],
            "placeholder_text_chunks": extraction_quality["placeholder_text_chunks"],
            "layout_elements": len(elements),
            "preserves_page_bbox": True,
        },
        "modality_processing": {
            "status": "ready",
            "text_chunks": sum(chunk.element_type == "text" for chunk in chunks),
            "circuit_diagrams": metadata["circuit_diagrams"],
            "formula_elements": metadata["formula_elements"],
            "formula_processing": formula_processing,
            "table_elements": metadata["table_elements"],
            "circuit_vision_model": metadata["vision_model"],
        },
        "knowledge_fusion": {
            "status": "ready",
            "vector_store": qdrant_status.get("mode", "faiss") if qdrant_status.get("enabled") else "faiss",
            "vector_points": len(chunks),
            "circuit_vector_points": qdrant_status.get("circuit_points", 0),
            "circuit_vector_store": (
                "qdrant+faiss"
                if qdrant_status.get("multimodal_qdrant_enabled")
                else "faiss"
                if qdrant_status.get("local_faiss_enabled")
                else "disabled"
            ),
            "graph_nodes": len(graph["nodes"]),
            "graph_edges": len(graph["edges"]),
            "graph_store": "neo4j" if neo4j_status.get("enabled") else "local-json",
        },
        "retrieval_service": {
            "status": "ready",
            "strategies": [
                "vector", "BM25", "knowledge-graph", "rerank", "circuit-image"
            ],
            "question_bank_search": False,
            "circuit_image_min_score": settings.circuit_image_retrieval_min_score,
            "circuit_image_max_references": settings.circuit_image_retrieval_max_references,
        },
        "application": {
            "status": "ready",
            "context_modalities": ["text", "formula", "table", "circuit-description", "netlist", "image"],
        },
    }
    (output_dir / "pipeline_audit.json").write_text(
        json.dumps(metadata["pipeline_layers"], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (output_dir / "index_meta.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    report(88, "artifacts_ready", "构建产物已生成，等待原子切换")
    logger.info("Knowledge base populated: %s", metadata)
    return metadata
